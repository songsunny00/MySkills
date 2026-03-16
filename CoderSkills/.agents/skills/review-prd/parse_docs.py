#!/usr/bin/env python3
"""解析产品需求文档目录，将 docx/xlsx 转为可读文本文件。

用法:
  python parse_docs.py <文档目录路径>            # 解析所有文件
  python parse_docs.py <文档目录路径> --latest    # 只解析每组文档的最新版本
"""

import sys
import os
import re
from pathlib import Path
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    print("需要安装 python-docx: pip install python-docx")
    sys.exit(1)

try:
    from openpyxl import load_workbook
except ImportError:
    print("需要安装 openpyxl: pip install openpyxl")
    sys.exit(1)


# 匹配文件名中的日期版本号，如 "PRD文档（1203）" 或 "字段规则（1122）" 或 "字段规则1014"
VERSION_PATTERN = re.compile(r"[（(]?(\d{4})[）)]?(?=\.\w+$|$)")


def extract_version_key(filename: str) -> tuple[str, str]:
    """从文件名提取基础名和版本号。

    返回 (base_name, version)，如果没有版本号则 version 为空字符串。
    例如: "样本中心需求PRD文档（1203）.docx" -> ("样本中心需求PRD文档.docx", "1203")
    """
    stem = Path(filename).stem
    suffix = Path(filename).suffix

    match = VERSION_PATTERN.search(stem)
    if match:
        version = match.group(1)
        # 移除版本号部分得到基础名
        base = stem[:match.start()].rstrip("（(） )") + suffix
        return base, version
    return stem + suffix, ""


def filter_latest_versions(filepaths: list[Path]) -> list[Path]:
    """对同名不同版本的文件只保留最新版本。"""
    groups: dict[str, list[tuple[str, Path]]] = defaultdict(list)

    for fp in filepaths:
        base, version = extract_version_key(fp.name)
        groups[base].append((version, fp))

    result = []
    for base, versions in groups.items():
        if len(versions) == 1:
            result.append(versions[0][1])
        else:
            # 按版本号降序，取最大的
            versions.sort(key=lambda x: x[0], reverse=True)
            latest = versions[0]
            result.append(latest[1])
            skipped = [v[1].name for v in versions[1:]]
            print(f"[SKIP] {base}: 保留最新版 {latest[1].name}，跳过旧版 {skipped}")

    return sorted(result)


def parse_docx(filepath: Path) -> str:
    """解析 docx 文件，提取所有文本内容（含表格）。"""
    doc = Document(str(filepath))
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1]

        if tag == "p":
            for para in doc.paragraphs:
                if para._element is element:
                    style = para.style.name if para.style else ""
                    prefix = ""
                    if "Heading" in style:
                        level = style.replace("Heading", "").strip()
                        prefix = "#" * int(level) + " " if level.isdigit() else "## "
                    text = para.text.strip()
                    if text:
                        parts.append(prefix + text)
                    break

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is element:
                    parts.append(_parse_table(table))
                    break

    return "\n\n".join(parts)


def _parse_table(table) -> str:
    """将 docx 表格转为 markdown 格式。"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if len(rows) >= 1:
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)
    return "\n".join(rows)


def parse_xlsx(filepath: Path) -> str:
    """解析 xlsx 文件，提取所有 sheet 的内容。"""
    wb = load_workbook(str(filepath), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")

        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append("| " + " | ".join(cells) + " |")

        if rows:
            col_count = rows[0].count("|") - 1
            separator = "| " + " | ".join(["---"] * col_count) + " |"
            rows.insert(1, separator)
            parts.append("\n".join(rows))

    wb.close()
    return "\n\n".join(parts)


def make_unique_path(output_dir: Path, name: str, suffix: str) -> Path:
    """生成不冲突的输出文件路径。"""
    base_path = output_dir / f"{name}_{suffix}.txt"
    if not base_path.exists():
        return base_path
    counter = 2
    while True:
        candidate = output_dir / f"{name}_{suffix}_{counter}.txt"
        if not candidate.exists():
            return candidate
        counter += 1


def main():
    if len(sys.argv) < 2:
        print("用法: python parse_docs.py <文档目录路径> [--latest]")
        sys.exit(1)

    doc_dir = Path(sys.argv[1])
    latest_only = "--latest" in sys.argv

    if not doc_dir.is_dir():
        print(f"目录不存在: {doc_dir}")
        sys.exit(1)

    output_dir = doc_dir / "_parsed"
    output_dir.mkdir(exist_ok=True)

    # 收集所有待处理文件
    all_files = []
    for filepath in sorted(doc_dir.rglob("*")):
        if "_parsed" in filepath.parts or filepath.name.startswith("."):
            continue
        if filepath.suffix.lower() in (".docx", ".xlsx"):
            all_files.append(filepath)

    # 多版本去重
    if latest_only:
        all_files = filter_latest_versions(all_files)
        print(f"[INFO] --latest 模式：筛选后 {len(all_files)} 个文件待解析\n")

    parsed_files = []
    used_names: set[str] = set()

    for filepath in all_files:
        suffix_tag = filepath.suffix[1:].lower()
        output_path = make_unique_path(output_dir, filepath.stem, suffix_tag)
        relative = filepath.relative_to(doc_dir)

        try:
            if suffix_tag == "docx":
                content = parse_docx(filepath)
                output_path.write_text(content, encoding="utf-8")
                parsed_files.append(str(output_path))
                print(f"[OK] {relative} -> {output_path.name}")

            elif suffix_tag == "xlsx":
                content = parse_xlsx(filepath)
                output_path.write_text(content, encoding="utf-8")
                parsed_files.append(str(output_path))
                print(f"[OK] {relative} -> {output_path.name}")

        except Exception as e:
            print(f"[ERR] {relative}: {e}")

    print(f"\n解析完成，共处理 {len(parsed_files)} 个文件，输出到 {output_dir}")


if __name__ == "__main__":
    main()